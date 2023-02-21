def foc():
	import os
	import pygame
	import time
	import menu 
	import docx
	from docx.shared import Pt
	pygame.init()
	
	screen = pygame.display.set_mode( (900,900) )

	font = pygame.font.SysFont(None, 32)

	clock = pygame.time.Clock()

	start_time = pygame.time.get_ticks() 


	def find(x,y,t):
		return [int(x/t), int(y/t)]

	output=open('files\pupilname')
	s=output.readline()
	desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
	doc = docx.Document(f'{desktop}\{s}.docx')
	paused  = False
	running = True

	tests = [[112, 105, 117, 126, 102, 123] , [122, 127, 109, 119, 131 ,108] , [107, 115, 134, 124, 104, 116] , [132, 136, 101, 111, 135, 128] , [118, 129, 114, 130, 133, 120] , [103, 110, 121, 125, 113, 106]]
	wait = 101

	#coordinats = [[],[],[],[],[],[]]
	coordinats = [[[300, 300] ,[350, 300] , [400, 300] , [450, 300], [500, 300], [550, 300]], [[300, 350], [350, 350], [400, 350],[450, 350],[500, 350],[550, 350]], [[300, 400], [350, 400], [400, 400], [450, 400], [500, 400], [550, 400] ], [[300, 450], [350, 450], [400, 450], [450, 450], [500, 450], [550, 450]], [[300, 500],[350, 500],[400, 500],[450, 500],[500, 500],[550, 500]], [[300, 550],[350, 550],[400, 550],[450, 550],[500, 550],[550, 550]]]


	HEIGHT = 600
	WIDHT = 600


	TILE = 50

	[pygame.draw.line(screen, pygame.Color('white'), (x, 300), (x, HEIGHT)) for x in range(300, (WIDHT + TILE), TILE)]
	[pygame.draw.line(screen, pygame.Color('white'), (300, y), (WIDHT,y)) for y in range(300, (HEIGHT + TILE), TILE)]

	#text = font.render('Hello!', 1, (255,255,255))

	ind = [300,300]
	col = 0

	for i1 in range(6):
		

		for i3 in tests[i1]:
			text = font.render( str(i3), 1, (255,255,255))
			screen.blit(text,(ind[0], ind[1]))
			ind[0] += 50
			
		col += 1
		ind = [300,300]
		ind[1] += 50 * col

	while running:
		
		
		
		if not paused:
			counting_time = pygame.time.get_ticks() - start_time

			# change milliseconds into minutes, seconds, milliseconds
			counting_minutes = str(int(counting_time/60000)).zfill(2)
			counting_seconds = str( int((counting_time%60000)/1000) ).zfill(2)
			counting_millisecond = str(counting_time%1000).zfill(3)

			counting_string = "%s:%s" % (counting_minutes, counting_seconds)

			counting_text = font.render(str(counting_string), 1, (255,255,255))
			counting_rect = counting_text.get_rect(center = (30, 10))

		screen.fill( 'black', rect=(0,0,100,200/2) )
		screen.blit(counting_text, counting_rect)

		pygame.display.update()

		clock.tick(25)
		
		
		for event in pygame.event.get():
			
			if event.type == pygame.MOUSEMOTION :
				p = event.pos
				
			if event.type == pygame.MOUSEBUTTONDOWN :           
				if event.button == 1:
					fin = find(p[0], p[1], TILE)
					xf = fin[1] - 6 
					yf = fin[0] - 6
					if xf==3 and yf==1:
						running=False
						t=int(counting_minutes)*60+int(counting_seconds)
						focusrec=open('files\ocusrec','r').read()

						if 648/t>6:
							doc.add_paragraph().add_run('Уровень внимания высокий '+'\n').font.size=Pt(16)
						elif 648/t<=6 and 648/t>=4:
							doc.add_paragraph().add_run('Уровень внимания средний '+'\n').font.size=Pt(16)
						else:
							doc.add_paragraph().add_run('Уровень внимания низкий '+'\n').font.size=Pt(16)
							doc.add_paragraph(str(focusrec)+'\n\n')
						output.close()
						
								 
					if (xf > 5) or (xf < 0) or (yf > 5) or (yf < 0):
						pass
					
					else:

						
						c1 = coordinats[xf][yf][0]
						c2 = coordinats[xf][yf][1]
						
						
						real = tests[xf][yf]
						
						if wait != real:
							text = font.render( "Ошибка", 1, "red" )
							screen.blit(text,(395, 100)) 
						
											
						#print(c1, c2)
						
						else:
							pygame.draw.rect(screen, pygame.Color('black') , (395, 100, 100, 100))
							
							pygame.draw.rect(screen, pygame.Color('white') , (c1, c2, TILE -2, TILE -2))
							
							t = tests[xf][yf]
							text = font.render( str(t), 1, (0,0,0))
							screen.blit(text,(c1, c2))
							
							wait += 1
						
							if wait > 136:
								
								paused = True
								
								screen.fill('blue')
																		
								text = font.render( 'Вы справились за ' + str(counting_string), 1, (0,0,0))
								screen.blit(text,(450, 450))
								pygame.draw.rect(screen, pygame.Color('blue') , (10, 10, 100, 100))


			
			if event.type == pygame.QUIT:
				running = False

			if event.type == pygame.KEYDOWN:
				if event.key == pygame.K_ESCAPE:
					running = False
						 
		

		
	doc.save(f'{desktop}\{s}.docx')
	pygame.quit()   
	output.close()
	menu.menu()

