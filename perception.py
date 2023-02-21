from tkinter import *
from time import *
import docx
from docx.shared import Pt
import os

def compute():
	import menu
	output=open('files\pupilname')
	s=output.readline()
	output.close()
	desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
	doc = docx.Document(f'{desktop}\{s}.docx')
	s1=first.get()
	s2=second.get()
	if s1=='4' and s2=='3':
		end=time()
		t=end-start
		percrec=open('files\percrec','r').read()
		if t<=20:
			doc.add_paragraph().add_run('Уровень восприятия очень высокий'+'\n').font.size=Pt(16)
		elif t>= 21 and t<=30:
			doc.add_paragraph().add_run('Уровень восприятия высокий'+'\n').font.size=Pt(16)
		elif t>=31 and t<=50:
			doc.add_paragraph().add_run('Уровень восприятия средний'+'\n').font.size=Pt(16)
		elif t>=51 and t<=60:
			doc.add_paragraph().add_run('Уровень восприятия низкий'+'\n').font.size=Pt(16)
			doc.add_paragraph(str(percrec)+'\n\n')
		else:
			doc.add_paragraph().add_run('Уровень восприятия очень низкий'+'\n').font.size=Pt(16)
			doc.add_paragraph(str(percrec)+'\n\n')
		doc.save(f'{desktop}\{s}.docx')
		perc.destroy()
		menu.menu()

	else:
		fail=Tk()
		fail.title("Ошибка")
		screen_width = fail.winfo_screenwidth()
		screen_height = fail.winfo_screenheight()
		x_cordinate = int((screen_width/2) - (300/2))
		y_cordinate = int((screen_height/2) - (150/2))
		fail.geometry("{}x{}+{}+{}".format(300,150,x_cordinate,y_cordinate))
		fail.resizable(False, False)
		fail.attributes('-toolwindow', True)
		rules=Label(fail, text='Неверный ответ', font=("Arial",18))
		rules.place(x=60,y=40)
		butclose = Button(fail, font=("Arial",10), text="Ввести заново", bg='White', fg='Black',command=lambda: (fail.destroy()))
		butclose.place(x=80,y=80, width=150, height=30)
		fail.bind('<Return>', lambda event:(fail.destroy() ))
		fail.mainloop()
def percep():
	global start,rule,first,second,perc,url,label
	start=time()
	perc=Tk()
	perc.title("Диагностика восприятия")
	screen_width = perc.winfo_screenwidth()
	screen_height = perc.winfo_screenheight()
	x_cordinate = int((screen_width/2) - (445/2))
	y_cordinate = int((screen_height/2) - (450/2))
	perc.geometry("{}x{}+{}+{}".format(445,450,x_cordinate,y_cordinate))
	perc.resizable(False, False)
	perc.attributes('-toolwindow', True)
	canvas = Canvas(perc,width = 500,height = 500)      
	canvas.pack()
	
	img = PhotoImage(file='files\perc.png') 
	first=Entry()
	first.place(x=35,y=300,width=150, height=30)
	second=Entry()
	second.place(x=255,y=300,width=150, height=30)     
	canvas.create_image(10,10,anchor=NW,image=img) 
	rule=Label(perc, text='Подберите такие узоры кусочков, которые более всего \n подходят к рисункам ковриков и впишите их номера', font=("Arial",10))
	rule.place(x=50,y=330)
	butmem = Button(perc, font=("Arial",12), text="Отправить ответ", bg='White', fg='Black',command=lambda: ( compute()))
	butmem.place(x=145,y=380, width=150, height=30)
	perc.bind('<Return>', lambda event:(compute()) )
	perc.mainloop()
