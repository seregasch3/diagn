from tkinter import *
import docx
from docx.shared import Pt
import os

def stri():
	string=ans.get("1.0","end").split()
	output=open('files\pupilname')
	s=output.readline()
	desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
	doc = docx.Document(f'{desktop}\{s}.docx')
	badmind='солнце,шнурки,лось,пол,горячий,очки,сани,волга,град,кастрюля,роза,помидор'.split(',')
	k=0
	for i in string:
		if i.lower() in badmind:
			k+=1
	if k<=5:
		mindrec=open('files\mindrec','r').read()
		doc.add_paragraph().add_run('Уровень мышления низкий '+'\n').font.size=Pt(16)
		doc.add_paragraph(str(mindrec)+'\n\n')
	elif k>=6 and k<=8:
		doc.add_paragraph().add_run('Уровень мышления средний '+'\n').font.size=Pt(16)
	elif k>=9 and k<=12:
		doc.add_paragraph().add_run('Уровень мышления высокий '+'\n').font.size=Pt(16)
	doc.save(f'{desktop}\{s}.docx')
	minder.destroy()
	output.close()
def mind():
	import menu
	global ans, minder
	mind=['лампа,фонарь,солнце,свеча','сапоги,ботинки,шнурки,валенки', 'собака,лошадь,корова,лось','стол,стул,пол,кровать','сладкий,горький,кислый,горячий','очки,глаза,нос,уши','трактор,комбайн,машина,сани','Москва,Киев,Волга,Минск','шум,свист,гром,град','суп,кисель,кастрюля,картошка','береза,сосна,дуб,роза','абрикос,персик,помидор,апельсин']
	minder = Tk()
	minder.title("Диагностика мышления")
	screen_width = minder.winfo_screenwidth()
	screen_height = minder.winfo_screenheight()
	x_cordinate = int((screen_width/2) - (500/2))
	y_cordinate = int((screen_height/2) - (580/2))
	minder.geometry("{}x{}+{}+{}".format(500,580,x_cordinate,y_cordinate))
	minder.resizable(False, False)
	minder.attributes('-toolwindow', True)
	rule=Label(minder, text='Выберите лишнее слово в каждом ряду и напечатайте его:', font=("Arial",13))
	rule.pack()
	words0=Label(minder, text=str(mind[0]), font=("Arial",18))
	words0.place(x=50,y=40,width=400)
	words1=Label(minder, text=str(mind[1]), font=("Arial",18))
	words1.place(x=50,y=70,width=400)
	words2=Label(minder, text=str(mind[2]), font=("Arial",18))
	words2.place(x=50,y=100,width=400)
	words3=Label(minder, text=str(mind[3]), font=("Arial",18))
	words3.place(x=50,y=130,width=400)
	words4=Label(minder, text=str(mind[4]), font=("Arial",18))
	words4.place(x=50,y=160,width=400)
	words5=Label(minder, text=str(mind[5]), font=("Arial",18))
	words5.place(x=50,y=190,width=400)
	words6=Label(minder, text=str(mind[6]), font=("Arial",18))
	words6.place(x=50,y=220,width=400)
	words7=Label(minder, text=str(mind[7]), font=("Arial",18))
	words7.place(x=50,y=250,width=400)
	words8=Label(minder, text=str(mind[8]), font=("Arial",18))
	words8.place(x=50,y=280,width=400)
	words9=Label(minder, text=str(mind[9]), font=("Arial",18))
	words9.place(x=50,y=310,width=400)
	words10=Label(minder, text=str(mind[10]), font=("Arial",18))
	words10.place(x=50,y=340,width=400)
	words11=Label(minder, text=str(mind[11]), font=("Arial",18))
	words11.place(x=50,y=370,width=400)
	ans=Text(minder,width=60, height=5)
	ans.place(x=10,y=430)
	butans=Button(minder, text="Отправить ответ",command=lambda:(stri(),menu.menu()))
	butans.place(x=200,y=530)
	minder.mainloop()
