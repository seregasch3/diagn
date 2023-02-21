from tkinter import *
import docx
from docx.shared import Pt
import os

memwords=['сено, ключ, самолет, поезд, картина\n','месяц, певец, радио, трава, перевал\n','автомобиль, сердце, букет, тротуар, столетие\n','фильм, аромат, горы, океан, неподвижность\n','календарь, мужчина, женщина, абстракция, вертолет\n']
def memo():
	mem=Tk()
	mem.title("Диагностика памяти")
	screen_width = mem.winfo_screenwidth()
	screen_height = mem.winfo_screenheight()
	x_cordinate = int((screen_width/2) - (900/2))
	y_cordinate = int((screen_height/2) - (400/2))
	mem.geometry("{}x{}+{}+{}".format(900,400,x_cordinate,y_cordinate))
	mem.resizable(False, False)
	mem.attributes('-toolwindow', True)
	word=Label(mem, text='У Вас есть одна минута, чтобы запомнить эти слова\n', font=("Arial",16))
	word.pack()
	words0=Label(mem, text=str(memwords[0]), font=("Arial",22))
	words0.pack()
	words1=Label(mem, text=str(memwords[1]), font=("Arial",22))
	words1.pack()
	words2=Label(mem, text=str(memwords[2]), font=("Arial",22))
	words2.pack()
	words3=Label(mem, text=str(memwords[3]), font=("Arial",22))
	words3.pack()
	words4=Label(mem, text=str(memwords[4]), font=("Arial",22))
	words4.pack()
	mem.after(60000,lambda: (mem.destroy(),memo2()))
	mem.mainloop()
	
def stri1():
	memrec=open('files\memrec','r').read()
	memrec
	string=ans.get("1.0","end").split()
	output=open('files\pupilname')
	s=output.readline()
	desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
	doc = docx.Document(f'{desktop}\{s}.docx')
	k=0
	for i in string:
		for j in memwords:
			if i.lower() in j:
				k+=1
	if k<=6:
		doc.add_paragraph().add_run('Уровень памяти низкий'+'\n').font.size=Pt(16)
		doc.add_paragraph(str(memrec)+'\n\n')
	elif k>=7 and k<=12:
		doc.add_paragraph().add_run('Уровень памяти ниже среднего '+'\n').font.size=Pt(16)
		doc.add_paragraph(str(memrec)+'\n\n')
	elif k>=13 and k<=17:
		doc.add_paragraph().add_run('Уровень памяти отличный '+'\n').font.size=Pt(16)
	else:
		doc.add_paragraph().add_run('Уровень памяти феноменальный '+'\n').font.size=Pt(16)
	doc.save(f'{desktop}\{s}.docx')
	mem2.destroy()
	output.close()
def memo2():
	import menu
	global ans, mem2
	mem2=Tk()
	mem2.title("Диагностика памяти")
	screen_width = mem2.winfo_screenwidth()
	screen_height = mem2.winfo_screenheight()
	x_cordinate = int((screen_width/2) - (500/2))
	y_cordinate = int((screen_height/2) - (500/2))
	mem2.geometry("{}x{}+{}+{}".format(500,500,x_cordinate,y_cordinate))
	mem2.resizable(False, False)
	mem2.attributes('-toolwindow', True)
	rule=Label(mem2, text='Введите слова, которые запомнили', font=("Arial",12))
	rule.pack()
	ans=Text(mem2,width=100)
	ans.pack()
	butans=Button(mem2, text="Отправить ответ",command= lambda:(stri1(),menu.menu()))
	butans.pack()
	mem2.mainloop()
