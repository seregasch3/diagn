
from tkinter import *
from menu import menu
import docx
from docx.shared import Pt
import os

def initial():
	s1=surname.get()
	s2=name.get()
	s3=classa.get()
	s=s1+' '+s2+' '+s3+''
	doc = docx.Document()
	doc.add_paragraph().add_run(s).font.size=Pt(18)
	desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
	doc.save(f'{desktop}\{s}.docx')
	open("files\pupilname",'w').write(s)
	butmain.after(10,main.destroy) 

main = Tk()
main.title("Регистрация")
screen_width = main.winfo_screenwidth()
screen_height = main.winfo_screenheight()
x_cordinate = int((screen_width/2) - (500/2))
y_cordinate = int((screen_height/2) - (250/2))
main.geometry("{}x{}+{}+{}".format(500,250,x_cordinate,y_cordinate))
main.resizable(False, False)
main.attributes('-toolwindow', True)


rule=Label(main, text='Диагностика познавательных процессов',font=("Arial",18))
rule.place(x=20,y=5)
rule1=Label(main, text='Фамилия',font=("Arial",18))
rule1.place(x=50,y=45)
surname=Entry()
surname.place(x=200,y=45,width=200, height=30)
rule2=Label(main, text='Имя',font=("Arial",18))
rule2.place(x=70,y=95)
name=Entry()
name.place(x=200,y=95,width=200, height=30)
rule3=Label(main, text='Класс',font=("Arial",18))
rule3.place(x=65,y=145)
classa=Entry()
classa.place(x=200,y=145,width=200, height=30)

butmain = Button(main, text="Начать", bg='White', fg='Black',
                              command=lambda: (initial(),menu()))
butmain.place(x=200,y=190, width=200, height=30)
main.bind('<Return>', lambda event:(initial(),menu()) )
                         
main.mainloop()
